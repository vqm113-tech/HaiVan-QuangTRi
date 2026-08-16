# ==========================================
# bulletin/bulletin_generator.py
# HẢI VĂN QUẢNG TRỊ 5.0 - DYNAMIC DOCX GENERATOR
#
# Khung bản tin (header/tiêu đề/bảng/footer) được dựng khớp theo mẫu bản tin
# chuẩn thực tế của Đài KTTV tỉnh Quảng Trị (QTRI_HVHV_20260424_1600.docx):
#   - Bảng 1/2/3: cột "Vùng biển dự báo" và cột nhóm yếu tố (Thủy triều/Sóng
#     biển/Dòng chảy) dùng GỘP Ô THẬT (rowspan) theo chiều dọc, không để ô
#     trống lặp lại như bản trước.
#   - Tiêu đề bảng không tô màu nền (mẫu chuẩn dùng nền trắng, chữ đậm).
#   - Cột "Yếu tố dự báo" có thêm dòng "Ngày" phía trên, đúng như mẫu.
# ==========================================

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


def set_cell_background(cell, hex_color):
    """Đặt màu nền cho cell trong bảng"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=100, bottom=100, left=100, right=100):
    """Thiết lập lề trong cho cell"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table):
    """Đặt đường viền mảnh chuẩn văn bản hành chính cho bảng"""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def format_cell(cell, text, bold=False, italic=False, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color=None):
    """Hàm bổ trợ định dạng ô trong bảng. `text` có thể chứa '\\n' để xuống dòng
    trong cùng 1 ô (ví dụ tiêu đề 2 dòng "Ngày" / "Yếu tố dự báo")."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lines = str(text).split('\n')
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(lines[0])
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    for extra_line in lines[1:]:
        p2 = cell.add_paragraph()
        p2.alignment = align
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(1)
        run2 = p2.add_run(extra_line)
        run2.bold = bold
        run2.italic = italic
        run2.font.size = Pt(size)
        run2.font.name = 'Times New Roman'
    set_cell_margins(cell)
    if bg_color:
        set_cell_background(cell, bg_color)


def merge_vertical(table, col_idx, row_start, row_end, text, bold=True, size=13):
    """Gộp các ô theo chiều dọc (rowspan) tại cột `col_idx`, từ hàng
    `row_start` đến `row_end` (bao gồm cả 2 đầu), rồi ghi `text` vào ô đã gộp,
    căn giữa theo chiều dọc — đúng kiểu trình bày bảng hành chính chuẩn thay
    vì để trống các ô còn lại như bản trước."""
    if row_end > row_start:
        merged = table.cell(row_start, col_idx).merge(table.cell(row_end, col_idx))
    else:
        merged = table.cell(row_start, col_idx)
    format_cell(merged, text, bold=bold, size=size)
    return merged


def set_cell_diagonal_border(cell):
    """Thêm đường chéo từ góc trên-trái xuống góc dưới-phải trong ô — kiểu
    tiêu đề bảng hành chính cổ điển, ngăn cách nhãn 'Ngày' (trên) và
    'Yếu tố dự báo' (dưới) đúng như bản tin mẫu."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:tl2br w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def format_diagonal_header_cell(cell, top_right_text, bottom_left_text, size=13):
    """Ô tiêu đề chia chéo kiểu bảng hành chính: `top_right_text` (ví dụ
    "Ngày") đặt căn phải ở nửa trên, `bottom_left_text` (ví dụ "Yếu tố dự
    báo") đặt căn trái ở nửa dưới, có đường chéo phân cách 2 nhãn."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.paragraph_format.space_before = Pt(1)
    p1.paragraph_format.space_after = Pt(1)
    r1 = p1.add_run(top_right_text)
    r1.bold = True
    r1.font.size = Pt(size)
    r1.font.name = 'Times New Roman'

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(1)
    r2 = p2.add_run(bottom_left_text)
    r2.bold = True
    r2.font.size = Pt(size)
    r2.font.name = 'Times New Roman'

    set_cell_margins(cell)
    set_cell_diagonal_border(cell)


def format_cell_vertical(cell, text, bold=True, italic=False, size=13, margins=(60, 60)):
    """Ô tiêu đề chữ XOAY DỌC (đọc từ dưới lên, kiểu `textDirection="btLr"`)
    — dùng cho 7 cột ngày ở Bảng 3, và cho các cột con Hx/Thời gian/Ngày/
    Hm/Thời gian/Ngày của Bảng 2 (bản tin mùa) & Phụ lục 1/2 (hồ sơ mùa),
    đúng như ảnh mẫu thật (đối chiếu XML gốc: mỗi cột con rộng ~440 dxa,
    cỡ chữ 12, in nghiêng, xoay dọc — quá hẹp để hiển thị ngang)."""
    set_cell_margins(cell, left=margins[0], right=margins[1])

    tcPr = cell._tc.get_or_add_tcPr()
    text_dir = OxmlElement('w:textDirection')
    text_dir.set(qn('w:val'), 'btLr')
    tcPr.append(text_dir)

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'


REGIONS = [
    ("Vùng biển ngoài khơi phía Bắc", "offshore_north"),
    ("Vùng biển ngoài khơi phía Nam", "offshore_south"),
    ("Vùng biển ven bờ phía Bắc", "coastal_north"),
    ("Vùng biển ven bờ phía Nam", "coastal_south"),
    ("Côn Cỏ", "con_co"),
]


def _add_letterhead(doc, data_dict, title_text, subtitle_text=None, title_size=13,
                     default_bulletin_num="HVHD/QTRI"):
    """Dựng phần quốc hiệu + số hiệu + tiêu đề, DÙNG CHUNG cho mọi loại bản
    tin (10 ngày / nguy hiểm / tháng / mùa) — cùng 1 khung hành chính KTTV,
    chỉ khác nội dung tiêu đề. Tách ra đây (thay vì lặp lại như 2 hàm gốc
    create_qtri_bulletin/create_qtri_warning_bulletin) để thêm bản tin mới
    không phải chép lại ~60 dòng dựng header mỗi lần.

    QUAN TRỌNG — khớp đúng cấu trúc bảng ở file mẫu thật (kiểm tra bằng
    python-docx trên QTRI_HV1T_20260801_1600.docx, QTRI_HVNH_...): đây là
    bảng 2 HÀNG x 2 CỘT (không phải 1 hàng + vertical_alignment=BOTTOM như
    trước đây). Hàng 1 chứa tên đơn vị (trái) và quốc hiệu (phải) — vì mỗi
    ô nằm trên MỘT HÀNG RIÊNG, cả 2 ô tự nhiên căn theo mép TRÊN của hàng đó
    (Word mặc định top-align), nên dòng đầu "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT
    NAM" thẳng hàng với dòng đầu "ĐÀI KHÍ TƯỢNG THỦY VĂN" — đúng yêu cầu
    "phải nằm ngang với nhau" — mà KHÔNG cần đếm dòng trống thủ công. Hàng 2
    chứa "Số: ..." (trái) và "Quảng Trị, ngày ..." (phải), cũng tự nhiên
    thẳng hàng vì cùng 1 hàng. "TỈNH QUẢNG TRỊ" / "Độc lập - Tự do - Hạnh
    phúc" in đậm VÀ gạch chân — theo đúng yêu cầu người dùng (ảnh mẫu thực
    tế đơn vị dùng có gạch chân 2 dòng này; file mẫu thô kiểm bằng
    python-docx trước đó không thấy `<w:u>` nhưng người dùng xác nhận bản
    thật có gạch chân, có thể do khác phiên bản/đơn vị soạn — ưu tiên theo
    xác nhận trực tiếp của người dùng).

    Độ rộng 2 cột TÍNH THEO LỀ TRANG THỰC TẾ của section hiện tại (không
    dùng số inch cố định như bản trước) — vì mỗi loại bản tin (10 ngày /
    nguy hiểm / tháng / mùa) nay có lề khác nhau đúng theo file mẫu thật,
    số inch cố định cũ sẽ tràn lề ở những bản tin có lề rộng hơn (HV1T,
    HVHM). Tỉ lệ cột trái/phải giữ nguyên theo mẫu gốc (~37% / 63%)."""
    content_width = (
        doc.sections[-1].page_width
        - doc.sections[-1].left_margin
        - doc.sections[-1].right_margin
    )
    col0_w = int(content_width * 0.369)
    col1_w = content_width - col0_w

    header_table = doc.add_table(rows=2, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = col0_w
    header_table.columns[1].width = col1_w
    for row in header_table.rows:
        row.cells[0].width = col0_w
        row.cells[1].width = col1_w

    # ---- Hàng 1: tên đơn vị (trái) | quốc hiệu (phải) ----
    cell_left = header_table.cell(0, 0)
    p_l1 = cell_left.paragraphs[0]
    p_l1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l1.paragraph_format.space_after = Pt(0)
    r_l1 = p_l1.add_run("ĐÀI KHÍ TƯỢNG THỦY VĂN")
    r_l1.font.size = Pt(12)

    p_l2 = cell_left.add_paragraph()
    p_l2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l2.paragraph_format.space_after = Pt(0)
    r_l2 = p_l2.add_run("TRUNG BỘ")
    r_l2.font.size = Pt(12)

    p_l3 = cell_left.add_paragraph()
    p_l3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l3.paragraph_format.space_after = Pt(0)
    r_l3 = p_l3.add_run("ĐÀI KHÍ TƯỢNG THỦY VĂN")
    r_l3.font.size = Pt(12)
    r_l3.bold = True

    p_l4 = cell_left.add_paragraph()
    p_l4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l4.paragraph_format.space_after = Pt(0)
    r_l4 = p_l4.add_run("TỈNH QUẢNG TRỊ")
    r_l4.font.size = Pt(12)
    r_l4.bold = True
    r_l4.underline = True

    cell_right = header_table.cell(0, 1)
    p_r1 = cell_right.paragraphs[0]
    p_r1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r1.paragraph_format.space_after = Pt(0)
    r_r1 = p_r1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r_r1.font.size = Pt(12)
    r_r1.bold = True

    p_r2 = cell_right.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r2.paragraph_format.space_after = Pt(0)
    r_r2 = p_r2.add_run("Độc lập - Tự do - Hạnh phúc")
    r_r2.font.size = Pt(12)
    r_r2.bold = True
    r_r2.underline = True

    # ---- Hàng 2: Số hiệu (trái) | Quảng Trị, ngày ... (phải) ----
    cell_left2 = header_table.cell(1, 0)
    p_l5 = cell_left2.paragraphs[0]
    p_l5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l5.paragraph_format.space_before = Pt(6)
    p_l5.paragraph_format.space_after = Pt(0)
    r_num = p_l5.add_run(f"Số: {data_dict.get('bulletin_num', default_bulletin_num)}")
    r_num.font.size = Pt(12)

    cell_right2 = header_table.cell(1, 1)
    p_r3 = cell_right2.paragraphs[0]
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r3.paragraph_format.space_before = Pt(6)
    p_r3.paragraph_format.space_after = Pt(0)
    issue_date = data_dict.get('issue_date', datetime.now().strftime('ngày %d tháng %m năm %Y'))
    r_date = p_r3.add_run(f"Quảng Trị, {issue_date}")
    r_date.font.size = Pt(12)
    r_date.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run(title_text)
    r_title.bold = True
    r_title.font.size = Pt(title_size)

    if subtitle_text:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(10)
        r_sub = p_sub.add_run(f"({subtitle_text})")
        r_sub.italic = True
        r_sub.font.size = Pt(title_size)


def _add_footer_meta_and_signature(doc, data_dict, recipients_text=None):
    """Dựng khối 'Thời gian ban hành tiếp theo / Tin phát lúc / Dự báo viên'
    + bảng 'Nơi nhận' / chữ ký lãnh đạo — DÙNG CHUNG cho mọi loại bản tin."""
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(12)
    r_meta = p_meta.add_run(
        f"Thời gian ban hành bản tin tiếp theo: {data_dict.get('next_issue_time', '')}.\n"
        f"Tin phát lúc: {data_dict.get('issue_time', '16h00')}.\n"
        f"Dự báo viên: {data_dict.get('forecasters', '')}."
    )
    r_meta.italic = True

    # Độ rộng 2 cột tính theo lề trang thực tế (xem giải thích ở
    # _add_letterhead) — giữ nguyên tỉ lệ mẫu gốc "Nơi nhận" / chữ ký (~48%/52%).
    content_width = (
        doc.sections[-1].page_width
        - doc.sections[-1].left_margin
        - doc.sections[-1].right_margin
    )
    col0_w = int(content_width * 0.478)
    col1_w = content_width - col0_w

    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False

    cell_f_left = footer_table.cell(0, 0)
    cell_f_left.width = col0_w
    p_rec = cell_f_left.paragraphs[0]
    p_rec.paragraph_format.space_after = Pt(0)
    p_rec.add_run("Nơi nhận:\n").bold = True

    recipients = recipients_text or (
        "- Văn phòng tỉnh ủy;\n"
        "- Văn phòng UBND tỉnh;\n"
        "- BCH PTDS tỉnh;\n"
        "- Sở NN&MT tỉnh;\n"
        "- Báo và Đài PTTH tỉnh;\n"
        "- Phòng QLDB & TT, DL KTTV (Cục KTTV);\n"
        "- Trung tâm TT&DL KTTV (Cục KTTV);\n"
        "- Phòng Dự báo KTTV (Đài KTTV Trung Bộ);\n"
        "- Các trạm KTTV, radar;\n"
        "- Lưu Đài tỉnh."
    )
    r_rec_body = p_rec.add_run(recipients)
    r_rec_body.font.size = Pt(11)
    r_rec_body.italic = True

    cell_f_right = footer_table.cell(0, 1)
    cell_f_right.width = col1_w
    p_sign = cell_f_right.paragraphs[0]
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sign.paragraph_format.space_after = Pt(0)

    r_sign_title = p_sign.add_run("KT. GIÁM ĐỐC\nPHÓ GIÁM ĐỐC\n\n\n\n")
    r_sign_title.bold = True
    r_sign_title.font.size = Pt(13)

    r_sign_name = p_sign.add_run(data_dict.get('leader_name', 'Đàm Hữu Tuyến'))
    r_sign_name.bold = True
    r_sign_name.font.size = Pt(13)


def create_qtri_bulletin(data_dict, *args, output_path="Ban_tin_Hai_van_Quang_Tri.docx", **kwargs):
    """
    Sinh bản tin dự báo hải văn 10 ngày tỉnh Quảng Trị, đúng khung mẫu chuẩn
    của Đài KTTV tỉnh Quảng Trị.
    """
    if not isinstance(data_dict, dict):
        data_dict = {}

    if len(args) >= 1 and isinstance(args[0], str):
        data_dict['forecasters'] = args[0]
    if len(args) >= 2 and isinstance(args[1], str):
        data_dict['issue_time'] = args[1]

    doc = Document()

    # Khổ giấy A4 + lề trang đúng khớp file mẫu thật QTRI_HVHV_20260810_1600.docx
    # (đo bằng python-docx: pgSz 11906x16838 = A4; pgMar top=1134/right=851/
    # bottom=1134/left=1701 twip = Top 2cm, Bottom 2cm, Left 3cm, Right 1.5cm).
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    style.font.color.rgb = RGBColor(0, 0, 0)

    # -------------------------------------------------------------------------
    # HEADER QUỐC HIỆU + TIÊU ĐỀ (dùng chung _add_letterhead — xem docstring
    # hàm đó để biết vì sao đổi từ 1 hàng bảng + vertical_alignment=BOTTOM
    # sang 2 hàng bảng, khớp đúng cấu trúc file mẫu thật)
    # -------------------------------------------------------------------------
    if 'period_text' not in data_dict:
        data_dict['period_text'] = 'Từ ngày 24/4 đến ngày 03/5/2026'
    _add_letterhead(
        doc, data_dict,
        title_text="BẢN TIN DỰ BÁO, CẢNH BÁO HẢI VĂN THỜI HẠN ĐẾN 10 NGÀY\nVÙNG BIỂN TỈNH QUẢNG TRỊ",
        subtitle_text=data_dict.get('period_text'),
        default_bulletin_num="HVHN-114/QTRI",
    )

    # -------------------------------------------------------------------------
    # 1. TÌNH HÌNH HẢI VĂN 24 GIỜ QUA

    # -------------------------------------------------------------------------
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(2)
    p1.add_run("1. Tình hình hải văn trong 24 giờ qua").bold = True

    p1_body = doc.add_paragraph()
    p1_body.paragraph_format.space_after = Pt(4)
    p1_body.paragraph_format.first_line_indent = Cm(1)
    p1_body.add_run(data_dict.get('sec1_text', 'Vùng biển tỉnh Quảng Trị có gió Tây Bắc cấp 2-3, độ cao sóng tại trạm Hải văn Cồn Cỏ phổ biến từ 0.25 – 0.75 m, hướng Đông Nam. Biển bình thường.'))

    # -------------------------------------------------------------------------
    # 2. DỰ BÁO THỜI TIẾT BIỂN TRONG 3 NGÀY
    # -------------------------------------------------------------------------
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.add_run("2. Dự báo thời tiết biển trong 3 ngày").bold = True
    p2_body = doc.add_paragraph()
    p2_body.paragraph_format.space_after = Pt(4)
    p2_body.paragraph_format.first_line_indent = Cm(1)
    p2_body.add_run(data_dict.get('sec2_text', 'Vùng biển Quảng Trị có mưa rào và dông vài nơi đến rải rác, trong cơn dông có khả năng xuất hiện lốc xoáy và gió giật mạnh; gió Tây Bắc cấp 3-4, biển bình thường.'))

    # BẢNG 1
    p_t1 = doc.add_paragraph()
    p_t1.paragraph_format.space_after = Pt(2)
    p_t1.add_run("Bảng 1. Dự báo thời tiết biển tỉnh Quảng Trị").bold = True

    days_3 = data_dict.get('days_3', ['Ngày 1', 'Ngày 2', 'Ngày 3'])
    t1_data = data_dict.get('table1_data', {})

    t1 = doc.add_table(rows=1, cols=5)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1)

    # Header: Vùng biển dự báo | Ngày / Yếu tố dự báo | 3 ngày
    hdr = t1.rows[0].cells
    format_cell(hdr[0], "Vùng biển dự báo", bold=True)
    format_diagonal_header_cell(hdr[1], "Ngày", "Yếu tố dự báo")
    for idx, d_str in enumerate(days_3[:3]):
        format_cell(hdr[2 + idx], d_str, bold=True)

    metrics = [
        ("Hiện tượng thời tiết", "weather"),
        ("Tầm nhìn xa", "visibility"),
        ("Hướng, tốc độ gió", "wind"),
        ("Tình trạng biển", "sea_state")
    ]

    for reg_title, reg_key in REGIONS:
        reg_dict = t1_data.get(reg_key, {})
        row_start = len(t1.rows)
        for m_title, m_key in metrics:
            row_cells = t1.add_row().cells
            format_cell(row_cells[1], m_title)
            val_list = reg_dict.get(m_key, ["-", "-", "-"])
            for d_i in range(3):
                val = val_list[d_i] if d_i < len(val_list) else "-"
                format_cell(row_cells[2 + d_i], val)
        row_end = len(t1.rows) - 1
        merge_vertical(t1, 0, row_start, row_end, reg_title)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # -------------------------------------------------------------------------
    # 3. DỰ BÁO HẢI VĂN TRONG 3 NGÀY
    # -------------------------------------------------------------------------
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(4)
    p3.add_run("3. Dự báo hải văn trong 3 ngày\n").bold = True
    p3.add_run(data_dict.get('sec3_text', 'Trong 3 ngày tới, độ cao sóng vùng biển tỉnh Quảng Trị tại trạm Hải văn Cồn Cỏ phổ biến từ 0.25 – 1.25m, hướng Đông Bắc, biển bình thường.'))

    p_t2 = doc.add_paragraph()
    p_t2.paragraph_format.space_after = Pt(2)
    p_t2.add_run("Bảng 2. Dự báo hải văn tỉnh Quảng Trị 3 ngày tới").bold = True

    t2_data = data_dict.get('table2_data', {})
    t2 = doc.add_table(rows=1, cols=6)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2)

    hdr2 = t2.rows[0].cells
    format_cell(hdr2[0], "Vùng biển dự báo", bold=True)
    # Gộp ngang cột "Yếu tố dự báo" với cột "Chỉ tiêu" (chỉ ở hàng tiêu đề —
    # 2 hàng dữ liệu bên dưới vẫn tách riêng: nhóm yếu tố | chỉ tiêu cụ thể)
    # để ô chia chéo đủ rộng, đúng như bản tin mẫu.
    hdr2_merged = hdr2[1].merge(hdr2[2])
    format_diagonal_header_cell(hdr2_merged, "Ngày", "Yếu tố dự báo")
    for idx, d_str in enumerate(days_3[:3]):
        format_cell(hdr2[3 + idx], d_str, bold=True)

    # Mẫu cấu hình dòng Bảng 2: (tên nhóm, nhãn hàng, khóa dữ liệu)
    sub_metrics = [
        ("Thủy triều", "Hx (m)", "tide_hx"),
        ("Thủy triều", "Thời gian", "tide_hx_time"),
        ("Thủy triều", "Hm (m)", "tide_hm"),
        ("Thủy triều", "Thời gian", "tide_hm_time"),
        ("Sóng biển", "H (m)", "wave_height"),
        ("Sóng biển", "Hướng", "wave_dir"),
        ("Dòng chảy", "Vận tốc (m/s)", "current_speed"),
        ("Dòng chảy", "Hướng", "current_dir"),
    ]

    for reg_title, reg_key in REGIONS:
        reg_dict = t2_data.get(reg_key, {})
        region_row_start = len(t2.rows)
        group_row_start = {}
        for group_name, sub_title, sub_key in sub_metrics:
            row_cells = t2.add_row().cells
            row_idx = len(t2.rows) - 1
            group_row_start.setdefault(group_name, row_idx)
            format_cell(row_cells[2], sub_title)

            val_list = reg_dict.get(sub_key, ["-", "-", "-"])
            for d_i in range(3):
                val = val_list[d_i] if d_i < len(val_list) else "-"
                format_cell(row_cells[3 + d_i], val)
        region_row_end = len(t2.rows) - 1

        # Gộp cột "Vùng biển dự báo" theo chiều dọc cho cả vùng (8 hàng)
        merge_vertical(t2, 0, region_row_start, region_row_end, reg_title)

        # Gộp cột nhóm yếu tố (Thủy triều/Sóng biển/Dòng chảy) theo từng nhóm
        group_names_in_order = []
        for group_name, _, _ in sub_metrics:
            if group_name not in group_names_in_order:
                group_names_in_order.append(group_name)
        for i, group_name in enumerate(group_names_in_order):
            g_start = group_row_start[group_name]
            g_end = (group_row_start[group_names_in_order[i + 1]] - 1
                     if i + 1 < len(group_names_in_order) else region_row_end)
            merge_vertical(t2, 1, g_start, g_end, group_name)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # -------------------------------------------------------------------------
    # 4. DỰ BÁO HẢI VĂN TỪ NGÀY THỨ 4 ĐẾN NGÀY THỨ 10
    # -------------------------------------------------------------------------
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(4)
    p4.add_run("4. Dự báo hải văn từ ngày thứ 4 đến ngày thứ 10\n").bold = True
    p4.add_run(data_dict.get('sec4_text', 'Từ ngày thứ 4 đến ngày thứ 10, mực nước triều có xu thế cao dần...'))

    p_t3 = doc.add_paragraph()
    p_t3.paragraph_format.space_after = Pt(2)
    p_t3.add_run("Bảng 3. Dự báo hải văn vùng biển tỉnh Quảng Trị từ ngày thứ 4 đến ngày thứ 10").bold = True

    days_7 = data_dict.get('days_7', [f"Ngày {i}" for i in range(4, 11)])
    t3_data = data_dict.get('table3_data', {})

    t3 = doc.add_table(rows=1, cols=10)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t3)

    hdr3 = t3.rows[0].cells
    format_cell(hdr3[0], "Vùng biển dự báo", bold=True)
    hdr3_merged = hdr3[1].merge(hdr3[2])
    format_diagonal_header_cell(hdr3_merged, "Ngày", "Yếu tố dự báo")
    for idx, d_str in enumerate(days_7[:7]):
        format_cell_vertical(hdr3[3 + idx], d_str, bold=True)

    sub_metrics_t3 = [
        ("Thủy triều", "Hx (m)", "tide_hx"),
        ("Thủy triều", "Thời gian", "tide_hx_time"),
        ("Thủy triều", "Hm (m)", "tide_hm"),
        ("Thủy triều", "Thời gian", "tide_hm_time"),
        ("Sóng biển", "H (m)", "wave_height"),
        ("Sóng biển", "Hướng", "wave_dir"),
    ]

    for reg_title, reg_key in REGIONS:
        reg_dict = t3_data.get(reg_key, {})
        region_row_start = len(t3.rows)
        group_row_start = {}
        for group_name, sub_title, sub_key in sub_metrics_t3:
            row_cells = t3.add_row().cells
            row_idx = len(t3.rows) - 1
            group_row_start.setdefault(group_name, row_idx)
            format_cell(row_cells[2], sub_title)

            val_list = reg_dict.get(sub_key, ["-"] * 7)
            for d_i in range(7):
                val = val_list[d_i] if d_i < len(val_list) else "-"
                format_cell(row_cells[3 + d_i], val)
        region_row_end = len(t3.rows) - 1

        merge_vertical(t3, 0, region_row_start, region_row_end, reg_title)

        group_names_in_order = []
        for group_name, _, _ in sub_metrics_t3:
            if group_name not in group_names_in_order:
                group_names_in_order.append(group_name)
        for i, group_name in enumerate(group_names_in_order):
            g_start = group_row_start[group_name]
            g_end = (group_row_start[group_names_in_order[i + 1]] - 1
                     if i + 1 < len(group_names_in_order) else region_row_end)
            merge_vertical(t3, 1, g_start, g_end, group_name)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # -------------------------------------------------------------------------
    # 5 & 6. CẢNH BÁO NGUY HIỂM & TÁC ĐỘNG
    # -------------------------------------------------------------------------
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(4)
    p5.add_run("5. Khả năng xuất hiện các hiện tượng thời tiết, hải văn nguy hiểm\n").bold = True
    p5.add_run(data_dict.get('sec5_text', 'Cần đề phòng các hiện tượng lốc xoáy và gió giật mạnh trong mưa dông.'))

    p6 = doc.add_paragraph()
    p6.paragraph_format.space_after = Pt(6)
    p6.add_run("6. Khả năng tác động đến môi trường, điều kiện sống, cơ sở hạ tầng, các hoạt động kinh tế - xã hội.\n").bold = True
    p6.add_run(data_dict.get('sec6_text', 'Tất cả các tàu, thuyền hoạt động trên vùng biển có nguy cơ chịu tác động của lốc xoáy và gió giật mạnh.'))

    # FOOTER META
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(12)
    r_meta = p_meta.add_run(
        f"Thời gian ban hành bản tin tiếp theo: {data_dict.get('next_issue_time', '16h00 ngày tiếp theo')}.\n"
        f"Tin phát lúc: {data_dict.get('issue_time', '16h00')}.\n"
        f"Dự báo viên: {data_dict.get('forecasters', 'Đàm Hữu Tuyến, Nguyễn Quang Hiếu')}."
    )
    r_meta.italic = True

    # -------------------------------------------------------------------------
    # NƠI NHẬN & CHỮ KÝ LÃNH ĐẠO
    # -------------------------------------------------------------------------
    # Độ rộng 2 cột tính theo lề trang thực tế (như _add_footer_meta_and_signature)
    content_width = (
        doc.sections[-1].page_width
        - doc.sections[-1].left_margin
        - doc.sections[-1].right_margin
    )
    col0_w = int(content_width * 0.478)
    col1_w = content_width - col0_w

    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False

    cell_f_left = footer_table.cell(0, 0)
    cell_f_left.width = col0_w
    p_rec = cell_f_left.paragraphs[0]
    p_rec.paragraph_format.space_after = Pt(0)
    p_rec.add_run("Nơi nhận:\n").bold = True

    recipients = (
        "- VP BCH PTDS tỉnh;\n"
        "- Báo & Đài PTTH Quảng Trị;\n"
        "- Trung tâm Dự báo KTTV quốc gia;\n"
        "- Phòng Quản lý DB & TT, DL KTTV (Cục KTTV);\n"
        "- Trung tâm TT&DL KTTV (Cục KTTV);\n"
        "- Phòng Dự báo (Đài Trung Bộ);\n"
        "- Các trạm KTTV, Ra đa trong tỉnh;\n"
        "- Lưu: Đài tỉnh."
    )
    r_rec_body = p_rec.add_run(recipients)
    r_rec_body.font.size = Pt(11)
    r_rec_body.italic = True

    cell_f_right = footer_table.cell(0, 1)
    cell_f_right.width = col1_w
    p_sign = cell_f_right.paragraphs[0]
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sign.paragraph_format.space_after = Pt(0)

    r_sign_title = p_sign.add_run("KT. GIÁM ĐỐC\nPHÓ GIÁM ĐỐC\n\n\n\n")
    r_sign_title.bold = True
    r_sign_title.font.size = Pt(13)

    r_sign_name = p_sign.add_run(data_dict.get('leader_name', 'Đàm Hữu Tuyến'))
    r_sign_name.bold = True
    r_sign_name.font.size = Pt(13)

    doc.save(output_path)
    return output_path


generate_bulletin = create_qtri_bulletin


def create_qtri_warning_bulletin(data_dict, forecaster=None, issue_time=None,
                                  output_path="Tin_Hai_Van_Nguy_Hiem_Quang_Tri.docx"):
    """
    Sinh "TIN DỰ BÁO GIÓ MẠNH, SÓNG LỚN TRÊN VÙNG BIỂN TỈNH QUẢNG TRỊ" —
    bản tin hải văn nguy hiểm (khác bản tin 10 ngày thường), đúng khung mẫu
    QTRI_HVNH_20260310_1600.docx. Dữ liệu đầu vào lấy từ
    bulletin.warning_data.build_warning_data().
    """
    if not isinstance(data_dict, dict):
        data_dict = {}
    if forecaster:
        data_dict['forecasters'] = forecaster
    if issue_time:
        data_dict['issue_time'] = issue_time

    doc = Document()
    # Khổ giấy A4 + lề trang đúng khớp file mẫu thật QTRI_HVNH_20260310_1600.docx
    # (đo bằng python-docx: pgSz 11910x16850 ~ A4; pgMar top=0/right=850/
    # bottom=280/left=1559 twip = Top 0cm, Bottom 0.5cm, Left 2.75cm,
    # Right 1.5cm — bản tin nhanh cần lề hẹp để vừa đúng 1 trang).
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.1)
        section.left_margin = Cm(2.75)
        section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    style.font.color.rgb = RGBColor(0, 0, 0)

    # HEADER QUỐC HIỆU (giống bản tin 10 ngày) — chiều rộng cột tính theo lề
    # trang thực tế (xem giải thích ở _add_letterhead), tránh tràn lề vì
    # HVNH dùng lề hẹp hơn các bản tin khác.
    content_width = (
        doc.sections[-1].page_width
        - doc.sections[-1].left_margin
        - doc.sections[-1].right_margin
    )
    col0_w = int(content_width * 0.369)
    col1_w = content_width - col0_w

    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False

    cell_left = header_table.cell(0, 0)
    cell_left.width = col0_w
    p_l1 = cell_left.paragraphs[0]
    p_l1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l1.paragraph_format.space_after = Pt(0)

    r_l1 = p_l1.add_run("ĐÀI KHÍ TƯỢNG THỦY VĂN\nTRUNG BỘ\n")
    r_l1.font.size = Pt(12)
    r_l1.bold = False

    r_l2 = p_l1.add_run("ĐÀI KHÍ TƯỢNG THỦY VĂN\n")
    r_l2.font.size = Pt(12)
    r_l2.bold = True

    r_l2b = p_l1.add_run("TỈNH QUẢNG TRỊ")
    r_l2b.font.size = Pt(12)
    r_l2b.bold = True
    r_l2b.underline = True

    p_l2 = cell_left.add_paragraph()
    p_l2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l2.paragraph_format.space_after = Pt(0)
    r_num = p_l2.add_run(f"\nSố: {data_dict.get('bulletin_num', 'HVNH-14/16h00/QTRI')}")
    r_num.font.size = Pt(12)

    cell_right = header_table.cell(0, 1)
    cell_right.width = col1_w
    p_r1 = cell_right.paragraphs[0]
    p_r1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r1.paragraph_format.space_after = Pt(0)

    r_r1 = p_r1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    r_r1.font.size = Pt(12)
    r_r1.bold = True
    r_r2 = p_r1.add_run("Độc lập - Tự do - Hạnh phúc")
    r_r2.font.size = Pt(12)
    r_r2.bold = True
    r_r2.underline = True

    p_r2 = cell_right.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r2.paragraph_format.space_after = Pt(0)
    issue_date = data_dict.get('issue_date', datetime.now().strftime('ngày %d tháng %m năm %Y'))
    r_date = p_r2.add_run(f"\nQuảng Trị, {issue_date}")
    r_date.font.size = Pt(12)
    r_date.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # TIÊU ĐỀ
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("TIN DỰ BÁO GIÓ MẠNH, SÓNG LỚN\nTRÊN VÙNG BIỂN TỈNH QUẢNG TRỊ")
    r_title.bold = True
    r_title.font.size = Pt(14)

    # 1. HIỆN TRẠNG ĐÃ QUA
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(2)
    p1.add_run("1. Hiện trạng đã qua\n").bold = True
    p1.add_run(data_dict.get('past_text', ''))

    # 2. DỰ BÁO DIỄN BIẾN TRONG 24 GIỜ TỚI
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.add_run("2. Dự báo diễn biến trong 24 giờ tới\n").bold = True
    p2.add_run(data_dict.get('next24h_text', ''))

    zone_rows = data_dict.get('zone_rows', [])
    t = doc.add_table(rows=2, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    hdr0 = t.rows[0].cells
    hdr1 = t.rows[1].cells
    merge_vertical(t, 0, 0, 1, "Thời điểm dự báo")
    merge_vertical(t, 1, 0, 1, "Vùng biển ảnh hưởng")
    merge_vertical(t, 2, 0, 1, "Gió mạnh (cấp Bô-pho)")
    hdr0_wave_merged = hdr0[3].merge(hdr0[4])
    format_cell(hdr0_wave_merged, "Độ cao sóng", bold=True)
    format_cell(hdr1[3], "Độ cao (mét)", bold=True)
    format_cell(hdr1[4], "Hướng", bold=True)

    row_start = len(t.rows)
    for zr in zone_rows:
        cells = t.add_row().cells
        format_cell(cells[1], zr.get('zone', ''), bold=True)
        format_cell(cells[2], zr.get('wind_text', ''))
        format_cell(cells[3], zr.get('wave_range', ''))
        format_cell(cells[4], zr.get('wave_dir', ''))
    row_end = len(t.rows) - 1
    merge_vertical(t, 0, row_start, row_end, data_dict.get('time_range_text', ''))

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # 3. CẢNH BÁO
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(2)
    p3.add_run("3. Cảnh báo\n").bold = True
    p3.add_run(data_dict.get('warning_text', ''))

    # 4. CẢNH BÁO CẤP ĐỘ RỦI RO THIÊN TAI
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(2)
    p4.add_run("4. Cảnh báo cấp độ rủi ro thiên tai trên biển: ").bold = True
    p4.add_run(data_dict.get('risk_level', 'Chưa đến mức cảnh báo'))

    # 5. DỰ BÁO TÁC ĐỘNG
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(2)
    p5.add_run("5. Dự báo tác động\n").bold = True
    p5.add_run(data_dict.get('impact_text', ''))

    # FOOTER META
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(3)
    r_meta = p_meta.add_run(
        f"Thời gian ban hành bản tin tiếp theo: {data_dict.get('next_issue_time', '')}.\n"
        f"Tin phát lúc: {data_dict.get('issue_time', '16h00')}.\n"
        f"Dự báo viên: {data_dict.get('forecasters', '')}."
    )
    r_meta.italic = True

    # NƠI NHẬN & CHỮ KÝ — chiều rộng cột tính theo lề trang thực tế
    col0_w = int(content_width * 0.478)
    col1_w = content_width - col0_w

    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False

    cell_f_left = footer_table.cell(0, 0)
    cell_f_left.width = col0_w
    p_rec = cell_f_left.paragraphs[0]
    p_rec.paragraph_format.space_after = Pt(0)
    p_rec.add_run("Nơi nhận:\n").bold = True
    recipients = (
        "- Văn phòng tỉnh ủy;\n"
        "- Văn phòng UBND tỉnh;\n"
        "- BCH PTDS tỉnh;\n"
        "- Sở NN&MT tỉnh;\n"
        "- Báo và Đài PTTH tỉnh;\n"
        "- Phòng QL DB&TT, DL KTTV (Cục KTTV);\n"
        "- Trung tâm TT&DL KTTV (Cục KTTV);\n"
        "- Phòng Dự báo KTTV (Đài KTTV Trung Bộ);\n"
        "- Các trạm KTTV, Radar;\n"
        "- Lưu Đài tỉnh."
    )
    r_rec_body = p_rec.add_run(recipients)
    r_rec_body.font.size = Pt(11)
    r_rec_body.italic = True

    cell_f_right = footer_table.cell(0, 1)
    cell_f_right.width = col1_w
    p_sign = cell_f_right.paragraphs[0]
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sign.paragraph_format.space_after = Pt(0)
    r_sign_title = p_sign.add_run("KT. GIÁM ĐỐC\nPHÓ GIÁM ĐỐC\n\n\n\n")
    r_sign_title.bold = True
    r_sign_title.font.size = Pt(13)
    r_sign_name = p_sign.add_run(data_dict.get('leader_name', 'Đàm Hữu Tuyến'))
    r_sign_name.bold = True
    r_sign_name.font.size = Pt(13)

    doc.save(output_path)
    return output_path


def _add_monthly_zone_table(doc, data_dict, title_prefix="Bảng 1: Dự báo vùng biển Quảng Trị"):
    """Dựng Bảng 1 (thủy triều 3 kỳ/tháng + sóng biển) của bản tin THÁNG.
    Tách riêng thành hàm dùng chung để hồ sơ dự báo (HS_) trang 2 có thể
    NHÚNG LẠI đúng bảng này — đúng như file mẫu thật HS_QTRI_HV1T_... (trang
    "Phần ghi thảo luận dự báo" lặp lại y nguyên Bảng 1 của bản tin chính,
    không phải một bảng khác)."""
    this_label = data_dict.get('this_month_label', '')
    p_t1 = doc.add_paragraph()
    p_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t1.paragraph_format.space_after = Pt(2)
    p_t1.add_run(f"{title_prefix} {this_label}").bold = True

    period_labels = data_dict.get('period_labels', ['01-10', '11-20', '21-30'])
    table_data = data_dict.get('table_data', {})
    regions = data_dict.get('regions', REGIONS)

    t1 = doc.add_table(rows=1, cols=6)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1)

    hdr = t1.rows[0].cells
    format_cell(hdr[0], "Vùng biển\ndự báo", bold=True)
    hdr_merged = hdr[1].merge(hdr[2])
    format_diagonal_header_cell(hdr_merged, "Ngày", "Yếu tố dự báo")
    for idx, p_label in enumerate(period_labels[:3]):
        format_cell(hdr[3 + idx], p_label, bold=True)

    sub_metrics = [
        ("Thủy triều", "Hx (m)", "tide_hx"),
        ("Thủy triều", "Thời gian", "tide_hx_time"),
        ("Thủy triều", "Hm (m)", "tide_hm"),
        ("Thủy triều", "Thời gian", "tide_hm_time"),
        ("Sóng biển", "H (m)", "wave_height"),
    ]

    for reg_title, reg_key in regions:
        reg_dict = table_data.get(reg_key, {})
        region_row_start = len(t1.rows)
        group_row_start = {}
        for group_name, sub_title, sub_key in sub_metrics:
            row_cells = t1.add_row().cells
            row_idx = len(t1.rows) - 1
            group_row_start.setdefault(group_name, row_idx)
            format_cell(row_cells[2], sub_title)

            val_list = reg_dict.get(sub_key, ["-", "-", "-"])
            for p_i in range(3):
                val = val_list[p_i] if p_i < len(val_list) else "-"
                format_cell(row_cells[3 + p_i], val)
        region_row_end = len(t1.rows) - 1

        merge_vertical(t1, 0, region_row_start, region_row_end, reg_title)

        group_names_in_order = []
        for group_name, _, _ in sub_metrics:
            if group_name not in group_names_in_order:
                group_names_in_order.append(group_name)
        for i, group_name in enumerate(group_names_in_order):
            g_start = group_row_start[group_name]
            g_end = (group_row_start[group_names_in_order[i + 1]] - 1
                     if i + 1 < len(group_names_in_order) else region_row_end)
            merge_vertical(t1, 1, g_start, g_end, group_name)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def create_qtri_monthly_bulletin(data_dict, forecaster=None, issue_time=None,
                                  output_path="Ban_tin_Hai_van_Thang_Quang_Tri.docx"):
    """
    Sinh "BẢN TIN DỰ BÁO, CẢNH BÁO HẢI VĂN THỜI HẠN THÁNG TỈNH QUẢNG TRỊ"
    (HV1T), đúng khung mẫu QTRI_HV1T_20260801_1600.docx. Dữ liệu đầu vào lấy
    từ bulletin.monthly_data.build_monthly_data().
    """
    if not isinstance(data_dict, dict):
        data_dict = {}
    if forecaster:
        data_dict['forecasters'] = forecaster
    if issue_time:
        data_dict['issue_time'] = issue_time

    doc = Document()
    # Khổ giấy A4 + lề trang đúng khớp file mẫu thật QTRI_HV1T_20260801_1600.docx
    # (đo bằng python-docx: pgSz 11907x16840 = A4; pgMar top=1134/right=1134/
    # bottom=1134/left=1701 twip = Top 2cm, Bottom 2cm, Left 3cm, Right 2cm).
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    style.font.color.rgb = RGBColor(0, 0, 0)

    _add_letterhead(
        doc, data_dict,
        title_text="BẢN TIN DỰ BÁO, CẢNH BÁO HẢI VĂN THỜI HẠN THÁNG\nTỈNH QUẢNG TRỊ",
        subtitle_text=data_dict.get('title_period', 'Tháng .. năm ....'),
        default_bulletin_num="HVHD-01/QTRI",
    )

    # -------------------------------------------------------------------------
    # 1. PHÂN TÍCH, ĐÁNH GIÁ HẢI VĂN NỔI BẬT THÁNG TRƯỚC
    # -------------------------------------------------------------------------
    prev_label = data_dict.get('prev_month_label', 'tháng trước')
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(2)
    p1.add_run(f"1. Phân tích, đánh giá hải văn nổi bật {prev_label}").bold = True

    p1a = doc.add_paragraph()
    p1a.paragraph_format.space_after = Pt(2)
    p1a.add_run("Sóng biển").bold = True
    p1b = doc.add_paragraph()
    p1b.paragraph_format.space_after = Pt(6)
    p1b.paragraph_format.first_line_indent = Cm(1)
    p1b.add_run(data_dict.get('sec1_wave_text', ''))

    p1c = doc.add_paragraph()
    p1c.paragraph_format.space_after = Pt(2)
    p1c.add_run("Triều cường").bold = True
    p1d = doc.add_paragraph()
    p1d.paragraph_format.space_after = Pt(6)
    p1d.paragraph_format.first_line_indent = Cm(1)
    p1d.add_run(data_dict.get('sec1_tide_text', ''))

    # -------------------------------------------------------------------------
    # 2. DỰ BÁO HẢI VĂN THÁNG NÀY
    # -------------------------------------------------------------------------
    this_label = data_dict.get('this_month_label', 'tháng này')
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.add_run(f"2. Dự báo hải văn {this_label}").bold = True

    p2a = doc.add_paragraph()
    p2a.paragraph_format.space_after = Pt(2)
    p2a.add_run("Sóng biển").bold = True
    p2b = doc.add_paragraph()
    p2b.paragraph_format.space_after = Pt(6)
    p2b.paragraph_format.first_line_indent = Cm(1)
    p2b.add_run(data_dict.get('sec2_wave_text', ''))

    p2c = doc.add_paragraph()
    p2c.paragraph_format.space_after = Pt(2)
    p2c.add_run("Triều cường").bold = True
    p2d = doc.add_paragraph()
    p2d.paragraph_format.space_after = Pt(8)
    p2d.paragraph_format.first_line_indent = Cm(1)
    p2d.add_run(data_dict.get('sec2_tide_text', ''))

    # -------------------------------------------------------------------------
    # BẢNG 1: DỰ BÁO VÙNG BIỂN QUẢNG TRỊ THEO 3 KỲ TRONG THÁNG
    # -------------------------------------------------------------------------
    _add_monthly_zone_table(doc, data_dict)

    # -------------------------------------------------------------------------
    # 3. HIỆN TƯỢNG NGUY HIỂM & 4. TÁC ĐỘNG
    # -------------------------------------------------------------------------
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(2)
    p3.add_run("3. Khả năng xuất hiện các hiện tượng thời tiết, hải văn nguy hiểm.").bold = True
    p3_body = doc.add_paragraph()
    p3_body.paragraph_format.space_after = Pt(6)
    p3_body.paragraph_format.first_line_indent = Cm(1)
    p3_body.add_run(data_dict.get('sec3_text', ''))

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(2)
    p4.add_run(
        "4. Khả năng tác động đến môi trường, điều kiện sống, cơ sở hạ tầng, "
        "các hoạt động kinh tế - xã hội."
    ).bold = True
    p4_body = doc.add_paragraph()
    p4_body.paragraph_format.space_after = Pt(10)
    p4_body.paragraph_format.first_line_indent = Cm(1)
    p4_body.add_run(data_dict.get('sec4_text', ''))

    _add_footer_meta_and_signature(doc, data_dict, recipients_text=(
        "- Văn phòng tỉnh ủy;\n"
        "- Văn phòng UBND tỉnh;\n"
        "- BCH PTDS tỉnh;\n"
        "- Sở NN&MT tỉnh;\n"
        "- Báo và Đài PTTH tỉnh;\n"
        "- Phòng QLDB & TT, DL KTTV (Cục KTTV);\n"
        "- Trung tâm TT&DL KTTV (Cục KTTV);\n"
        "- Phòng Dự báo KTTV (Đài KTTV Trung Bộ);\n"
        "- Các trạm KTTV, rada;\n"
        "- Lưu Đài tỉnh."
    ))

    doc.save(output_path)
    return output_path


def _format_time_wrap(t):
    """Chèn khoảng trắng sau 'h' (VD "07h30" -> "07h 30") để chuỗi giờ có
    thể tự xuống dòng gọn trong ô cột hẹp (~0.3in) — đúng cách trình bày
    của mẫu thật (XML gốc ghi "15h 00", "06h 00", không phải "15h00")."""
    if not t:
        return t
    if 'h' in t and ' ' not in t:
        return t.replace('h', 'h ', 1)
    return t


def _render_tide_zone_table(doc, data_dict, title, col_widths, use_cm=False):
    """Dựng PHẦN THÂN bảng thủy triều 3 tháng (Vị trí dự báo x Nước lớn/
    Nước ròng x Hx/Thời gian/Ngày/Hm/Thời gian/Ngày) — dùng chung cho cả
    'Bảng 2' của bản tin MÙA (đơn vị mét, theo yêu cầu quy đổi mét trước
    đó) lẫn 'Phụ lục 1/2' của hồ sơ dự báo MÙA (đơn vị cm — ĐÚNG mẫu thật
    HS_QTRI_HVHM_..., có ghi chú riêng '(cm)' khác với bản tin chính).

    LUÔN ở khổ ĐỨNG (portrait) — đúng mẫu thật (đối chiếu bằng python-docx:
    toàn bộ tài liệu HS_QTRI_HVHM_20260615_1700.docx chỉ có 1 sectPr duy
    nhất, khổ 11907x16840 = A4 đứng, KHÔNG có trang ngang nào). Để vừa đủ
    18 cột con (3 tháng x 6 cột) trong khổ đứng, các cột con Hx/Thời gian/
    Ngày/Hm/Thời gian/Ngày XOAY DỌC 90° (textDirection btLr), cỡ chữ 12,
    cột rộng ~440 dxa (~0.3in) — sao chép chính xác từ mẫu thật."""
    forecast_months = data_dict.get('forecast_months', [])
    month_labels = [f"Tháng {mo}/{yr}" for yr, mo in forecast_months]
    n_months = len(forecast_months) or 3
    cols_per_month = 6

    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_after = Pt(2)
    p_t.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t.add_run(title).bold = True

    t2 = doc.add_table(rows=3, cols=1 + n_months * cols_per_month)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.autofit = False
    set_table_borders(t2)

    region_col_width, col_widths_list = col_widths
    total_cols = 1 + n_months * cols_per_month
    t2.columns[0].width = region_col_width
    for c_i in range(1, total_cols):
        t2.columns[c_i].width = col_widths_list[c_i - 1]
    for row in t2.rows:
        row.cells[0].width = region_col_width
        for c_i in range(1, total_cols):
            row.cells[c_i].width = col_widths_list[c_i - 1]

    merge_vertical(t2, 0, 0, 2, "Vị trí dự báo", size=12)

    unit_label = "(cm)" if use_cm else "(m)"
    for m_i, label in enumerate(month_labels):
        base = 1 + m_i * cols_per_month
        month_hdr = t2.cell(0, base).merge(t2.cell(0, base + cols_per_month - 1))
        format_cell(month_hdr, label, bold=True, italic=True, size=12)
        set_cell_margins(month_hdr, left=20, right=20)
        nl_hdr = t2.cell(1, base).merge(t2.cell(1, base + 2))
        format_cell(nl_hdr, "Nước lớn", bold=False, italic=True, size=12)
        set_cell_margins(nl_hdr, left=20, right=20)
        nr_hdr = t2.cell(1, base + 3).merge(t2.cell(1, base + 5))
        format_cell(nr_hdr, "Nước ròng", bold=False, italic=True, size=12)
        set_cell_margins(nr_hdr, left=20, right=20)
        sub_labels = [f"Hx {unit_label}", "Thời gian", "Ngày", f"Hm {unit_label}", "Thời gian", "Ngày"]
        row2 = t2.rows[2].cells
        for s_i, s_label in enumerate(sub_labels):
            format_cell_vertical(row2[base + s_i], s_label, bold=False, italic=True,
                                  size=12, margins=(20, 20))
    # Chữ xoay dọc cần CHIỀU CAO hàng đủ lớn mới hiển thị được (LibreOffice/
    # Word không tự tính đúng chiều cao cho ô xoay 90° nếu để mặc định) —
    # ép chiều cao tối thiểu ~1.35in cho hàng tiêu đề con này.
    t2.rows[2].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    t2.rows[2].height = Inches(1.35)

    table2_data = data_dict.get('table2_data', {})
    regions = data_dict.get('regions', REGIONS)
    for reg_title, reg_key in regions:
        row_cells = t2.add_row().cells
        row_cells[0].width = region_col_width
        format_cell(row_cells[0], reg_title, size=12)
        month_rows = table2_data.get(reg_key, [])
        for m_i in range(n_months):
            base = 1 + m_i * cols_per_month
            mrow = month_rows[m_i] if m_i < len(month_rows) else {}
            hx = mrow.get('hx')
            hx_t = mrow.get('hx_time')
            hx_d = mrow.get('hx_day')
            hm = mrow.get('hm')
            hm_t = mrow.get('hm_time')
            hm_d = mrow.get('hm_day')
            if use_cm:
                hx_str = "-" if hx is None else str(int(round(hx)))
                hm_str = "-" if hm is None else str(int(round(hm)))
            else:
                hx_str = "-" if hx is None else f"{hx / 100.0:.2f}"
                hm_str = "-" if hm is None else f"{hm / 100.0:.2f}"
            vals = [
                hx_str, _format_time_wrap(hx_t) or "-", "-" if hx_d is None else f"{hx_d:02d}",
                hm_str, _format_time_wrap(hm_t) or "-", "-" if hm_d is None else f"{hm_d:02d}",
            ]
            for s_i, v in enumerate(vals):
                row_cells[base + s_i].width = col_widths_list[base + s_i - 1]
                format_cell(row_cells[base + s_i], v, size=12)
                set_cell_margins(row_cells[base + s_i], left=20, right=20)

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_after = Pt(6)
    r_note = p_note.add_run(
        f"Ghi chú: Hx: độ cao thủy triều lớn nhất {unit_label}; Hm: độ cao thủy triều nhỏ nhất {unit_label}."
    )
    r_note.italic = True
    r_note.font.size = Pt(11)


def _compute_tide_table_col_widths(page_width, orig_l, orig_r, n_months):
    """Độ rộng cột ĐÚNG như mẫu thật (đối chiếu XML gốc
    HS_QTRI_HVHM_20260615_1700.docx): cột "Vị trí dự báo" ~1178 dxa
    (~0.82in), MỖI cột con trong 6 cột con (Hx/Thời gian/Ngày/Hm/Thời
    gian/Ngày) rộng ĐỀU NHAU ~440 dxa (~0.3in) — không lệch theo nội dung
    như bản tính tỉ lệ trước đó. `page_width` là chiều rộng trang ĐỨNG
    hiện tại (không hoán đổi w/h như khi còn dùng khổ ngang)."""
    usable_dxa = (page_width - orig_l - orig_r) / 635  # EMU -> dxa (20 dxa = 1pt = 12700 EMU)
    region_dxa = 1178
    subcol_dxa = 440
    total_dxa = region_dxa + n_months * 6 * subcol_dxa
    if total_dxa > usable_dxa:
        # An toàn: nếu số tháng nhiều hơn 3 khiến bảng vượt lề, co đều tỉ lệ
        # lại để vừa đúng khổ đứng, không để tràn lề.
        scale = usable_dxa / total_dxa
        region_dxa = int(region_dxa * scale)
        subcol_dxa = int(subcol_dxa * scale)
    region_col_width = Twips(region_dxa)
    col_widths_in = [Twips(subcol_dxa)] * (n_months * 6)
    return region_col_width, col_widths_in


def _add_seasonal_tide_table(doc, data_dict):
    """Dựng Bảng 2 (thủy triều 3 tháng tới, đủ 6 cột con Hx/Thời gian/Ngày/
    Hm/Thời gian/Ngày mỗi tháng — ĐÚNG cấu trúc mẫu thật) — LUÔN giữ khổ
    ĐỨNG (portrait) như phần còn lại của bản tin, KHÔNG chuyển sang khổ
    ngang (đối chiếu mẫu thật: toàn bộ tài liệu chỉ có 1 khổ đứng duy
    nhất). Các cột con xoay dọc 90° để vừa đủ chỗ — xem _render_tide_zone_table.

    Đơn vị MÉT — dùng cho BẢN TIN CHÍNH (đã quy đổi mét theo yêu cầu trước
    đó). Hồ sơ (HS_) dùng bảng Phụ lục riêng bằng cm — xem
    _add_seasonal_appendix_tables() bên dưới, KHÔNG dùng hàm này."""
    section = doc.sections[-1]
    n_months = len(data_dict.get('forecast_months', [])) or 3
    col_widths = _compute_tide_table_col_widths(
        section.page_width, section.left_margin, section.right_margin, n_months)
    title = f"Bảng 2: Dự báo thuỷ triều {data_dict.get('forecast_period_label', '')}"
    _render_tide_zone_table(doc, data_dict, title, col_widths, use_cm=False)


def _add_seasonal_appendix_tables(doc, data_dict):
    """Dựng 'Phụ lục 1: Kết quả dự báo theo phương pháp phân tích hàm điều
    hòa' + 'Phụ lục 2: Chọn kết quả dự báo' trong hồ sơ dự báo MÙA (HS_) —
    ĐÚNG cấu trúc mẫu thật (đơn vị CM, không phải mét như bản tin chính —
    đối chiếu bằng python-docx trên HS_QTRI_HVHM_20260615_1700.docx, ghi
    chú gốc ghi rõ '(cm)'). LUÔN giữ khổ ĐỨNG (portrait) — đối chiếu mẫu
    thật: toàn bộ tài liệu chỉ có 1 sectPr duy nhất, khổ A4 đứng, không hề
    có trang ngang nào; các cột con xoay dọc 90° để vừa đủ chỗ.

    Dự án chỉ có DUY NHẤT 1 phương pháp tính triều thật (mô hình điều hòa —
    bulletin/tide_model.py); mẫu thật có phương pháp thứ 2 ('mô hình số
    trị') nhưng dự án không có nguồn dữ liệu/mô hình số trị triều nào khác
    để so sánh, nên Phụ lục 2 ('Chọn kết quả') dùng LẠI đúng kết quả của
    Phụ lục 1 — trung thực với những gì dự án thực sự tính được, không bịa
    thêm một bộ số liệu giả từ "phương án khác" không tồn tại."""
    section = doc.sections[-1]
    n_months = len(data_dict.get('forecast_months', [])) or 3
    col_widths = _compute_tide_table_col_widths(
        section.page_width, section.left_margin, section.right_margin, n_months)

    _render_tide_zone_table(
        doc, data_dict, "Phụ lục 1: Kết quả dự báo theo phương pháp phân tích hàm điều hòa",
        col_widths, use_cm=True,
    )
    _render_tide_zone_table(
        doc, data_dict, "Phụ lục 2: Chọn kết quả dự báo",
        col_widths, use_cm=True,
    )


def create_qtri_seasonal_bulletin(data_dict, forecaster=None, issue_time=None,
                                   output_path="Ban_tin_Hai_van_Mua_Quang_Tri.docx"):
    """
    Sinh "BẢN TIN DỰ BÁO, CẢNH BÁO HẢI VĂN THỜI HẠN MÙA VÙNG BIỂN TỈNH QUẢNG
    TRỊ" (HVHM), đúng khung mẫu QTRI_HVHM_20260615_1700.docx. Dữ liệu đầu
    vào lấy từ bulletin.seasonal_data.build_seasonal_data().
    """
    if not isinstance(data_dict, dict):
        data_dict = {}
    if forecaster:
        data_dict['forecasters'] = forecaster
    if issue_time:
        data_dict['issue_time'] = issue_time

    doc = Document()
    # Khổ giấy A4 + lề trang đúng khớp file mẫu thật QTRI_HVHM_20260615_1700.docx
    # (đo bằng python-docx: pgSz 11907x16840 = A4; pgMar top=1134/right=1134/
    # bottom=1134/left=1701 twip = Top 2cm, Bottom 2cm, Left 3cm, Right 2cm —
    # giống hệt lề của bản tin HV1T).
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    style.font.color.rgb = RGBColor(0, 0, 0)

    _add_letterhead(
        doc, data_dict,
        title_text="BẢN TIN DỰ BÁO, CẢNH BÁO HẢI VĂN THỜI HẠN MÙA\nVÙNG BIỂN TỈNH QUẢNG TRỊ",
        subtitle_text=data_dict.get('title_period', 'Từ tháng .. đến tháng ..'),
        default_bulletin_num="HVHM-01/QTRI",
    )

    # -------------------------------------------------------------------------
    # 1. PHÂN TÍCH 02 THÁNG QUA
    # -------------------------------------------------------------------------
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(2)
    p1.add_run("1. Phân tích, đánh giá tình hình hải văn nổi bật trong 02 tháng qua").bold = True
    p1_body = doc.add_paragraph()
    p1_body.paragraph_format.space_after = Pt(6)
    p1_body.paragraph_format.first_line_indent = Cm(1)
    p1_body.add_run(data_dict.get('sec1_text', ''))

    station_name = data_dict.get('table1_station_name', 'Cồn Cỏ')
    table1_labels = data_dict.get('table1_labels', [])
    table1_rows = data_dict.get('table1_rows', [])

    p_t1 = doc.add_paragraph()
    p_t1.paragraph_format.space_after = Pt(2)
    p_t1.add_run(
        f"Bảng 1: Đặc trưng sóng, thủy triều tại trạm Hải văn {station_name}"
    ).bold = True

    t1 = doc.add_table(rows=2, cols=2 + len(table1_labels))
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1)
    hdr0 = t1.rows[0].cells
    hdr1 = t1.rows[1].cells
    merge_vertical(t1, 0, 0, 1, "Yếu tố")
    merge_vertical(t1, 1, 0, 1, "")
    for idx, label in enumerate(table1_labels):
        format_cell(hdr0[2 + idx], label, bold=True)
        t1.cell(1, 2 + idx).merge(t1.cell(1, 2 + idx))  # placeholder, filled below

    t1_rows_spec = [
        ("Thủy triều", "Nước lớn", "Hmax (m)", "hmax"),
        ("Thủy triều", "Nước lớn", "Ngày xuất hiện", "hmax_day"),
        ("Thủy triều", "Nước ròng", "Hmin (m)", "hmin"),
        ("Thủy triều", "Nước ròng", "Ngày xuất hiện", "hmin_day"),
        ("Sóng", "Độ cao sóng lớn nhất (m)", "", "wave_max"),
        ("Sóng", "Hướng sóng", "", "wave_dir"),
        ("Sóng", "Ngày xuất hiện", "", "wave_day"),
    ]

    group_row_start = {}
    for group_name, sub_title, sub_sub, key in t1_rows_spec:
        row_cells = t1.add_row().cells
        row_idx = len(t1.rows) - 1
        group_row_start.setdefault(group_name, row_idx)
        label = sub_title if not sub_sub else f"{sub_title}\n{sub_sub}"
        format_cell(row_cells[1], label)
        for m_i, ext in enumerate(table1_rows):
            val = ext.get(key) if ext else None
            if val is not None and key in ("hmax", "hmin"):
                val = f"{val / 100.0:.2f}"
            format_cell(row_cells[2 + m_i], "-" if val is None else str(val))
    region_row_end = len(t1.rows) - 1
    group_names_in_order = []
    for group_name, _, _, _ in t1_rows_spec:
        if group_name not in group_names_in_order:
            group_names_in_order.append(group_name)
    for i, group_name in enumerate(group_names_in_order):
        g_start = group_row_start[group_name]
        g_end = (group_row_start[group_names_in_order[i + 1]] - 1
                 if i + 1 < len(group_names_in_order) else region_row_end)
        merge_vertical(t1, 0, g_start, g_end, group_name)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -------------------------------------------------------------------------
    # 2. DỰ BÁO HẢI VĂN 3 THÁNG TỚI
    # -------------------------------------------------------------------------
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.add_run(f"2. Dự báo hải văn {data_dict.get('forecast_period_label', '')}").bold = True
    p2_body = doc.add_paragraph()
    p2_body.paragraph_format.space_after = Pt(6)
    p2_body.paragraph_format.first_line_indent = Cm(1)
    p2_body.add_run(data_dict.get('sec2_text', ''))

    p2w = doc.add_paragraph()
    p2w.paragraph_format.space_after = Pt(2)
    p2w.add_run("Cảnh báo khả năng xuất hiện các hiện tượng hải văn nguy hiểm:").bold = True
    p2w_body = doc.add_paragraph()
    p2w_body.paragraph_format.space_after = Pt(6)
    p2w_body.paragraph_format.first_line_indent = Cm(1)
    p2w_body.add_run(data_dict.get('sec2_warning_text', ''))

    p2i = doc.add_paragraph()
    p2i.paragraph_format.space_after = Pt(2)
    p2i.add_run(
        "Khả năng tác động đến môi trường, điều kiện sống, cơ sở hạ tầng, "
        "các hoạt động kinh tế - xã hội:"
    ).bold = True
    p2i_body = doc.add_paragraph()
    p2i_body.paragraph_format.space_after = Pt(6)
    p2i_body.paragraph_format.first_line_indent = Cm(1)
    p2i_body.add_run(data_dict.get('sec2_impact_text', ''))

    _add_seasonal_tide_table(doc, data_dict)

    # -------------------------------------------------------------------------
    # 3. XU THẾ HẢI VĂN 3 THÁNG SAU (CHỈ VĂN BẢN, KHÔNG CÓ BẢNG — đúng mẫu)
    # -------------------------------------------------------------------------
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(2)
    p3.add_run(f"3. Xu thế hải văn {data_dict.get('xu_the_period_label', '')}").bold = True
    p3_body = doc.add_paragraph()
    p3_body.paragraph_format.space_after = Pt(6)
    p3_body.paragraph_format.first_line_indent = Cm(1)
    p3_body.add_run(data_dict.get('sec3_text', ''))

    p3w = doc.add_paragraph()
    p3w.paragraph_format.space_after = Pt(2)
    p3w.add_run(
        "Cảnh báo khả năng xuất hiện triều cường và các hiện tượng hải văn nguy hiểm khác:"
    ).bold = True
    p3w_body = doc.add_paragraph()
    p3w_body.paragraph_format.space_after = Pt(6)
    p3w_body.paragraph_format.first_line_indent = Cm(1)
    p3w_body.add_run(data_dict.get('sec3_warning_text', ''))

    p3i = doc.add_paragraph()
    p3i.paragraph_format.space_after = Pt(2)
    p3i.add_run(
        "Khả năng tác động đến môi trường, điều kiện sống, cơ sở hạ tầng, "
        "các hoạt động kinh tế - xã hội:"
    ).bold = True
    p3i_body = doc.add_paragraph()
    p3i_body.paragraph_format.space_after = Pt(10)
    p3i_body.paragraph_format.first_line_indent = Cm(1)
    p3i_body.add_run(data_dict.get('sec3_impact_text', ''))

    _add_footer_meta_and_signature(doc, data_dict, recipients_text=(
        "- VP BCH PTDS tỉnh;\n"
        "- Báo & Đài PTTH Quảng Trị;\n"
        "- Trung tâm Dự báo KTTV quốc gia;\n"
        "- Phòng Quản lý DB & TT, DL KTTV (Cục KTTV);\n"
        "- Trung tâm TT&DL KTTV (Cục KTTV);\n"
        "- Phòng Dự báo (Đài Trung Bộ);\n"
        "- Các trạm KTTV, Ra đa trong tỉnh;\n"
        "- Lưu: Đài tỉnh."
    ))

    doc.save(output_path)
    return output_path
