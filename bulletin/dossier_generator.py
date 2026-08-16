# ==========================================
# bulletin/dossier_generator.py
# Sinh "HỒ SƠ DỰ BÁO ..." (file HS_QTRI_...) — hồ sơ nội bộ ghi lại quy
# trình thu thập số liệu / phân tích / thực hiện phương án / thảo luận dự
# báo, theo khung 8 mục CHUNG cho cả 3 loại bản tin có mẫu thật trong dự án
# (HVNH/HV1T/HVHM — xem HS_QTRI_*.docx người dùng cung cấp):
#   1. Thu thập, xử lý các loại thông tin dữ liệu
#   2. Phân tích đánh giá hiện trạng
#   3. Thực hiện các phương án dự báo
#   4. Thảo luận dự báo (nội dung chi tiết ở trang sau)
#   5. Xây dựng bản tin dự báo
#   6. Cung cấp bản tin dự báo
#   7. Bổ sung, cập nhật bản tin
#   8. Đánh giá chất lượng bản tin trước
# + trang 2: "Phần ghi thảo luận dự báo" / "THẢO LUẬN DỰ BÁO..." (nhắc lại
#   dự báo/cảnh báo chi tiết, NHÚNG LẠI đúng bảng vùng biển của bản tin
#   chính — không phải một bảng khác, đã kiểm tra bằng python-docx trên cả
#   3 file mẫu thật).
#
# ⚠️ QUAN TRỌNG — đối chiếu bằng python-docx trên cả 3 file mẫu thật phát
# hiện có 2 KIỂU trình bày mục 1-3 khác nhau tùy đơn vị/loại bản tin, không
# phải 1 khung chung như bản trước đây tự đặt ra:
#   - style="detailed" (khớp HS_QTRI_HVNH_...): bảng 4 CỘT, cột 0 ghi
#     "Hoàn thành trước giờ phát tin X'" (merge dọc theo cả mục), có thêm
#     hàng "Kết luận" tổng hợp cuối mỗi mục.
#   - style="simple" (khớp HS_QTRI_HV1T_..., HS_QTRI_HVHM_...): bảng 3 CỘT,
#     cột 0 ghi mã a/b/c/d (hoặc để trống), KHÔNG có hàng "Kết luận" riêng.
#
# ⚠️ Đây là hồ sơ NGHIỆP VỤ NỘI BỘ — phần lớn nội dung (nguồn số liệu tham
# khảo, phân tích synop, đánh giá chất lượng bản tin trước) đòi hỏi nhận
# định của dự báo viên, KHÔNG có mô hình nào trong dự án tạo ra được. Hàm
# này chỉ dựng ĐÚNG KHUNG bảng biểu hành chính và điền lại các nội dung đã
# tính được từ bản tin tương ứng (văn bản dự báo/cảnh báo, bảng vùng biển)
# — các trường còn lại để trống/placeholder, dự báo viên tự điền trên giao
# diện trước khi xuất.
# ==========================================

from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from bulletin.bulletin_generator import (
    set_table_borders, format_cell, merge_vertical, set_cell_margins,
    _add_monthly_zone_table, _add_seasonal_tide_table, _add_seasonal_appendix_tables,
)


def _add_row(table, col0_text, col1_text, col3_text, bold_col0=False):
    """Dùng cho style='detailed' (4 cột: mã thời gian | nhãn (gộp cột 1+2)
    | nội dung) — cột 0 chỉ điền ở hàng đầu mục rồi merge dọc sau.

    QUAN TRỌNG: đối chiếu file mẫu thật HS_QTRI_HVNH_... bằng python-docx —
    cột 1 và cột 2 LUÔN được gộp làm MỘT ô nhãn rộng hơn (không phải cột 2
    bỏ trống cạnh cột 1 như bản trước đây), nếu không sẽ để lại một cột hẹp
    toàn ô trống chạy dọc bảng — đúng lỗi "nhiều ô trống chưa gộp lại"."""
    cells = table.add_row().cells
    if col0_text:
        format_cell(cells[0], col0_text, bold=bold_col0, align=WD_ALIGN_PARAGRAPH.LEFT)
    label_cell = cells[1].merge(cells[2])
    format_cell(label_cell, col1_text or "", align=WD_ALIGN_PARAGRAPH.LEFT)
    format_cell(cells[3], col3_text or "", align=WD_ALIGN_PARAGRAPH.LEFT)
    return cells


def _add_conclusion_row(table, label_text, value_text):
    """Hàng 'Kết luận' — đối chiếu file mẫu thật: cột 0+1+2 gộp làm MỘT ô
    nhãn (không phải chỉ cột 0 như các hàng thường), cột 3 = nội dung."""
    cells = table.add_row().cells
    label_cell = cells[0].merge(cells[1]).merge(cells[2])
    format_cell(label_cell, label_text, align=WD_ALIGN_PARAGRAPH.LEFT)
    format_cell(cells[3], value_text or "", align=WD_ALIGN_PARAGRAPH.LEFT)
    return cells


def _add_section_header_simple(t, n_cols, title):
    cells = t.add_row().cells
    merged = cells[0]
    for c in cells[1:]:
        merged = merged.merge(c)
    format_cell(merged, title, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)


def _add_item_row_simple(t, code, label, value):
    cells = t.add_row().cells
    format_cell(cells[0], code or "", align=WD_ALIGN_PARAGRAPH.LEFT)
    format_cell(cells[1], label or "", align=WD_ALIGN_PARAGRAPH.LEFT)
    format_cell(cells[2], value or "", align=WD_ALIGN_PARAGRAPH.LEFT)
    return cells


def _add_tail_row_simple(t, label, value):
    cells = t.add_row().cells
    merged = cells[0].merge(cells[1])
    format_cell(merged, label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    format_cell(cells[2], value or "", align=WD_ALIGN_PARAGRAPH.LEFT)


def create_forecast_dossier(data_dict, output_path="Ho_so_du_bao_Quang_Tri.docx"):
    """
    data_dict keys DÙNG CHUNG cho cả 2 style:
      style ('detailed' | 'simple')     — mặc định 'detailed' (khớp HVNH)
      title, issue_time_text, unit_text, shift_leader, forecasters
      section5_file_ref, section6_text, section7_text, section8_text
      discussion_title, discussion_forecaster_note
      zone_table_kind (None | 'monthly' | 'seasonal' | 'warning')
      zone_table_data (dict)            — data_dict gốc của bản tin chính,
          truyền lại nguyên vẹn để trang 2 NHÚNG LẠI đúng Bảng 1/Bảng 2 thật
          (dùng _add_monthly_zone_table / _add_seasonal_tide_table) thay vì
          một bảng khác.

    style='detailed' (HVNH) — thêm:
      section1_rows/section2_rows/section3_rows: list[(label, value)]
      section1_conclusion/section2_conclusion/section3_conclusion
      discussion_intro, discussion_zone_table (list[dict] — bảng gió/sóng
          riêng cho tin nguy hiểm, dùng khi zone_table_kind=None)
      discussion_warning, discussion_risk, discussion_impact

    style='simple' (HV1T, HVHM) — thêm:
      section1_items/section2_items: list[(code, label, value)]
      section3_items: list[(code_or_None, label, value)] — code=None thì
          KHÔNG in lại ở cột 0 (dòng tiếp theo cùng nhóm, giống mẫu thật)
      section4_extra: list[(label, value)] | None — mục 4 con a/b/c (chỉ
          HV1T có, HVHM không) — nếu có sẽ chèn ngay dưới hàng mục 4 chính
      discussion_body: list[(heading, text)] — các đoạn "Sóng biển:",
          "Triều cường:"... hiển thị trước bảng vùng biển ở trang 2
    """
    if not isinstance(data_dict, dict):
        data_dict = {}
    style = data_dict.get('style', 'detailed')
    zone_kind = data_dict.get('zone_table_kind')

    doc = Document()
    # Khổ giấy A4 + lề trang đúng khớp từng file hồ sơ mẫu thật (đo bằng
    # python-docx trên cả 3 file HS_QTRI_*.docx người dùng cung cấp — mỗi
    # loại hồ sơ có lề khác nhau, không dùng chung 1 khung như bản trước):
    #   - style='simple', zone_table_kind='seasonal_appendix' (HVHM):
    #     HS_QTRI_HVHM_20260615_1700.docx -> Top 2cm, Bottom 2cm, Left 2.5cm, Right 2cm
    #   - style='simple', zone_table_kind='monthly' (HV1T):
    #     HS_QTRI_HV1T_20260801_1600.docx -> Top 1.5cm, Bottom 1.5cm, Left 2cm, Right 2cm
    #   - style='detailed' (HVNH, mặc định):
    #     HS_QTRI_HVNH_20260103_1600.docx -> Top 1.5cm, Bottom 1.5cm, Left 2cm, Right 1.5cm
    if style == 'simple' and zone_kind == 'seasonal_appendix':
        m_top, m_bottom, m_left, m_right = 2, 2, 2.5, 2
    elif style == 'simple':
        m_top, m_bottom, m_left, m_right = 1.5, 1.5, 2, 2
    else:
        m_top, m_bottom, m_left, m_right = 1.5, 1.5, 2, 1.5
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(m_top)
        section.bottom_margin = Cm(m_bottom)
        section.left_margin = Cm(m_left)
        section.right_margin = Cm(m_right)

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)

    # ---- TIÊU ĐỀ ----
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(0)
    r_title = p_title.add_run(data_dict.get('title', 'HỒ SƠ DỰ BÁO HẢI VĂN'))
    r_title.bold = True
    r_title.font.size = Pt(14)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(8)
    r_sub = p_sub.add_run("TỈNH QUẢNG TRỊ")
    r_sub.bold = True
    r_sub.font.size = Pt(14)

    for line in [
        f"Thời gian phát tin theo quy định: {data_dict.get('issue_time_text', '')}",
        f"Đơn vị dự báo: {data_dict.get('unit_text', 'Đài KTTV tỉnh Quảng Trị.')}",
        f"Trưởng ca dự báo: {data_dict.get('shift_leader', '')}",
        f"Các dự báo viên: {data_dict.get('forecasters', '')}",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(line)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    if style == 'simple':
        _build_simple_style_table(doc, data_dict)
    else:
        _build_detailed_style_table(doc, data_dict)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---- Chữ ký trưởng ca ----
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sign.add_run("Trưởng ca dự báo").bold = True
    p_sign2 = doc.add_paragraph()
    p_sign2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sign2.paragraph_format.space_before = Pt(30)
    p_sign2.add_run(data_dict.get('shift_leader', ''))

    # =========================================================================
    # TRANG 2: PHẦN GHI THẢO LUẬN DỰ BÁO — NHÚNG LẠI đúng bảng của bản tin
    # chính (đúng như cả 3 file mẫu thật đều làm), không phải bảng khác.
    # =========================================================================
    doc.add_page_break()

    p2_title = doc.add_paragraph()
    p2_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2_title.add_run(data_dict.get('discussion_title', data_dict.get('title', '')))
    r.bold = True
    r.font.size = Pt(14)

    p2_sub = doc.add_paragraph()
    p2_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2_sub.add_run("TỈNH QUẢNG TRỊ")
    r2.bold = True
    r2.font.size = Pt(14)

    p2_note = doc.add_paragraph()
    p2_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2_note.add_run("(Phần ghi thảo luận dự báo)").italic = True

    p2_leader = doc.add_paragraph()
    p2_leader.add_run(f"Trưởng ca dự báo: {data_dict.get('shift_leader', '')}")

    for heading, text in data_dict.get('discussion_body', []):
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(8)
        p_h.paragraph_format.space_after = Pt(2)
        p_h.add_run(heading).bold = True
        p_b = doc.add_paragraph()
        p_b.paragraph_format.first_line_indent = Cm(1)
        p_b.add_run(text or '')

    if not data_dict.get('discussion_body') and data_dict.get('discussion_intro'):
        p2_body_title = doc.add_paragraph()
        p2_body_title.paragraph_format.space_before = Pt(8)
        p2_body_title.add_run("Dự báo diễn biến:").bold = True
        p2_body = doc.add_paragraph()
        p2_body.paragraph_format.first_line_indent = Cm(1)
        p2_body.add_run(data_dict.get('discussion_intro', ''))

    zone_kind = data_dict.get('zone_table_kind')
    zone_data = data_dict.get('zone_table_data')
    if data_dict.get('zone_table_intro'):
        p_zi = doc.add_paragraph()
        p_zi.paragraph_format.space_before = Pt(6)
        p_zi.add_run(data_dict['zone_table_intro'])
    if zone_kind == 'monthly' and zone_data:
        _add_monthly_zone_table(doc, zone_data)
    elif zone_kind == 'seasonal' and zone_data:
        _add_seasonal_tide_table(doc, zone_data)
    elif zone_kind == 'seasonal_appendix' and zone_data:
        _add_seasonal_appendix_tables(doc, zone_data)
    else:
        zone_table = data_dict.get('discussion_zone_table')
        if zone_table:
            tz = doc.add_table(rows=2, cols=5)
            tz.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(tz)
            hz0 = tz.rows[0].cells
            hz1 = tz.rows[1].cells
            merge_vertical(tz, 0, 0, 1, "Thời điểm dự báo")
            merge_vertical(tz, 1, 0, 1, "Vùng biển ảnh hưởng")
            merge_vertical(tz, 2, 0, 1, "Gió mạnh (cấp Bô-pho)")
            hz0_wave = hz0[3].merge(hz0[4])
            format_cell(hz0_wave, "Độ cao sóng", bold=True)
            format_cell(hz1[3], "Độ cao (mét)", bold=True)
            format_cell(hz1[4], "Hướng", bold=True)
            row_start = len(tz.rows)
            for zr in zone_table:
                cells = tz.add_row().cells
                format_cell(cells[1], zr.get('zone', ''), bold=True)
                format_cell(cells[2], zr.get('wind', ''))
                format_cell(cells[3], zr.get('wave_range', ''))
                format_cell(cells[4], zr.get('wave_dir', ''))
            row_end = len(tz.rows) - 1
            merge_vertical(tz, 0, row_start, row_end, data_dict.get('discussion_time_range', ''))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if data_dict.get('discussion_warning'):
        p_w_title = doc.add_paragraph()
        p_w_title.paragraph_format.space_before = Pt(6)
        p_w_title.add_run("Cảnh báo").bold = True
        p_w_body = doc.add_paragraph()
        p_w_body.paragraph_format.first_line_indent = Cm(1)
        p_w_body.add_run(data_dict['discussion_warning'])

    if data_dict.get('discussion_risk'):
        p_risk = doc.add_paragraph()
        p_risk.add_run("Cảnh báo cấp độ rủi ro thiên tai trên biển: ").bold = True
        p_risk.add_run(data_dict['discussion_risk'])

    if data_dict.get('discussion_impact'):
        p_imp_title = doc.add_paragraph()
        p_imp_title.paragraph_format.space_before = Pt(6)
        p_imp_title.add_run("Dự báo tác động").bold = True
        p_imp_body = doc.add_paragraph()
        p_imp_body.paragraph_format.first_line_indent = Cm(1)
        p_imp_body.add_run(data_dict['discussion_impact'])

    if data_dict.get('discussion_forecaster_note'):
        p_fn = doc.add_paragraph()
        p_fn.paragraph_format.space_before = Pt(10)
        p_fn.add_run(data_dict['discussion_forecaster_note'])

    doc.save(output_path)
    return output_path


def _build_detailed_style_table(doc, data_dict):
    """style='detailed' — khớp HS_QTRI_HVNH_...: bảng 4 cột, cột 0 = 'Hoàn
    thành trước giờ phát tin X'' (merge dọc), có hàng 'Kết luận' cuối mục."""
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    t.columns[0].width = Inches(1.1)
    t.columns[1].width = Inches(1.0)
    t.columns[2].width = Inches(0.4)
    t.columns[3].width = Inches(4.3)
    hdr = t.rows[0].cells
    hdr_merged = hdr[0].merge(hdr[1]).merge(hdr[2]).merge(hdr[3])
    format_cell(hdr_merged, "1. Thu thập, xử lý các loại thông tin dữ liệu", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    s1_rows = data_dict.get('section1_rows') or [
        ("Các loại bản tin", "Tham khảo bản tin của TTDB Quốc Gia và Đài khu vực"),
        ("Số liệu viễn thám, quan trắc", ""),
        ("Các sản phẩm mô hình", "Nchmf, ecmwf, GMS, KMA"),
    ]
    row_start = len(t.rows)
    for label, value in s1_rows:
        _add_row(t, None, label, value)
    row_end = len(t.rows) - 1
    if row_end >= row_start:
        merge_vertical(t, 0, row_start, row_end, "Hoàn thành trước giờ phát tin 75'", bold=False)
    _add_conclusion_row(t, "Kết luận\n(tính đầy đủ, có bổ sung, chỉnh lý)",
                         data_dict.get('section1_conclusion', 'Đầy đủ'))

    hdr2 = t.add_row().cells
    hdr2_merged = hdr2[0].merge(hdr2[1]).merge(hdr2[2]).merge(hdr2[3])
    format_cell(hdr2_merged, "2. Phân tích đánh giá hiện trạng", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    s2_rows = data_dict.get('section2_rows') or [
        ("Hình thế chính (24 giờ qua)", ""),
        ("Diễn biến hải văn nguy hiểm", ""),
    ]
    row_start = len(t.rows)
    for label, value in s2_rows:
        _add_row(t, None, label, value)
    row_end = len(t.rows) - 1
    if row_end >= row_start:
        merge_vertical(t, 0, row_start, row_end, "Hoàn thành trước giờ phát tin 60'", bold=False)
    _add_conclusion_row(t, "Kết luận\n(thời gian bắt đầu và kết thúc,\nkhu vực, định lượng)",
                         data_dict.get('section2_conclusion', ''))

    hdr3 = t.add_row().cells
    hdr3_merged = hdr3[0].merge(hdr3[1]).merge(hdr3[2]).merge(hdr3[3])
    format_cell(hdr3_merged, "3. Thực hiện các phương án dự báo", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    s3_rows = data_dict.get('section3_rows') or [
        ("Phương pháp synop", data_dict.get('discussion_intro', '')),
        ("Phân tích số trị (hướng, thời gian và\nphạm vi ảnh hưởng, cường độ)", ""),
        ("Phân tích viễn thám", ""),
    ]
    row_start = len(t.rows)
    for label, value in s3_rows:
        _add_row(t, None, label, value)
    row_end = len(t.rows) - 1
    if row_end >= row_start:
        merge_vertical(t, 0, row_start, row_end, "Hoàn thành trước giờ phát tin 30'", bold=False)
    _add_conclusion_row(t, "Kết luận\n(thời gian bắt đầu và kết thúc,\nkhu vực, định lượng gió)",
                         data_dict.get('section3_conclusion', data_dict.get('discussion_intro', '')))

    tail_rows = [
        ("4. Thảo luận dự báo:\nHoàn thành trước giờ phát tin 15'", "Nội dung chi tiết ghi trang sau"),
        ("5. Xây dựng bản tin dự báo:\nHoàn thành trước giờ phát tin 5'",
         f"Bản tin đính kèm hồ sơ này\n{data_dict.get('section5_file_ref', '')}"),
        ("6. Cung cấp bản tin dự báo:\nFax, Email, cập nhật web và các trục trặc",
         data_dict.get('section6_text', '')),
        ("7. Bổ sung, cập nhật bản tin\n(Thời gian và các thông tin cập nhật)",
         data_dict.get('section7_text', 'Không cập nhật/ bổ sung')),
        ("8. Đánh giá chất lượng bản tin trước", data_dict.get('section8_text', '')),
    ]
    for label, value in tail_rows:
        cells = t.add_row().cells
        merged = cells[0].merge(cells[1]).merge(cells[2])
        format_cell(merged, label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell(cells[3], value, align=WD_ALIGN_PARAGRAPH.LEFT)


def _build_simple_style_table(doc, data_dict):
    """style='simple' — khớp HS_QTRI_HV1T_..., HS_QTRI_HVHM_...: bảng 3
    cột, cột 0 = mã a/b/c/d (hoặc trống), KHÔNG có hàng 'Kết luận' riêng."""
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    t.columns[0].width = Inches(1.1)
    t.columns[1].width = Inches(2.0)
    t.columns[2].width = Inches(3.4)
    for row in t.rows:
        row.cells[0].width = Inches(1.1)
        row.cells[1].width = Inches(2.0)
        row.cells[2].width = Inches(3.4)
    hdr = t.rows[0].cells
    hdr_merged = hdr[0].merge(hdr[1]).merge(hdr[2])
    format_cell(hdr_merged, "1. Thu thập, xử lý các loại thông tin dữ liệu", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    s1_items = data_dict.get('section1_items') or [
        ("a", "Thu thập số liệu quan trắc", ""),
        ("b", "Xử lý các loại thông tin dữ liệu", ""),
        ("c", "Cập nhật số liệu", ""),
    ]
    for code, label, value in s1_items:
        _add_item_row_simple(t, code, label, value)

    _add_section_header_simple(t, 3, "2. Phân tích đánh giá hiện trạng")
    s2_items = data_dict.get('section2_items') or [
        ("a", "Phân tích các dữ liệu quan trắc để xác định điều kiện hải văn đã qua và hiện tại", ""),
        ("b", "Phân tích diễn biến của yếu tố mực nước/sóng trên cơ sở dữ liệu quan trắc và sản phẩm mô hình", ""),
    ]
    for code, label, value in s2_items:
        _add_item_row_simple(t, code, label, value)

    _add_section_header_simple(t, 3, "3. Thực hiện các phương án dự báo")
    s3_items = data_dict.get('section3_items') or [
        (None, "Phương án 1", ""),
        (None, "Phương án 2", ""),
    ]
    for code, label, value in s3_items:
        _add_item_row_simple(t, code, label, value)

    # mục 4
    section4_extra = data_dict.get('section4_extra')
    _add_tail_row_simple(t, "4. Thảo luận dự báo:", "Nội dung chi tiết ghi trang sau")
    if section4_extra:
        _add_section_header_simple(t, 3, "a. Phân tích độ tin cậy các phương án:")
        for label, value in section4_extra:
            _add_item_row_simple(t, "-", label, value)
        _add_section_header_simple(
            t, 3, "b. Tổng hợp kết quả dự báo các phương án và ý kiến của dự báo viên (ghi ở trang sau)")
        _add_section_header_simple(t, 3, "c. Kết luận của người chủ trì (ghi ở trang sau)")

    _add_tail_row_simple(t, "5. Xây dựng bản tin dự báo:",
                          f"Bản tin đính kèm hồ sơ này: {data_dict.get('section5_file_ref', '')}")
    _add_tail_row_simple(t, "6. Cung cấp bản tin dự báo:", data_dict.get('section6_text', ''))
    _add_tail_row_simple(t, "7. Bổ sung, cập nhật bản tin:",
                          data_dict.get('section7_text', 'Không'))
    _add_tail_row_simple(t, "8. Đánh giá chất lượng bản tin trước:",
                          data_dict.get('section8_text', ''))

