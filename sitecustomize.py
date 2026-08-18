"""Streamlit Cloud compatibility and seasonal bulletin formatting fixes."""
import mimetypes
import os
import re
import sys
from pathlib import Path


# Streamlit Cloud/Linux: replace the legacy Windows folder-opening action
# with browser downloads.
if not hasattr(os, "startfile") and sys.platform != "win32":
    def _startfile_linux_compat(path):
        try:
            import streamlit as st
            output_dir = Path(path)
            st.markdown("### 📥 TẢI FILE BÁO CÁO")
            if not output_dir.exists():
                st.info("Chưa có file báo cáo để tải.")
                return
            files = sorted(
                (p for p in output_dir.iterdir() if p.is_file()),
                key=lambda p: p.stat().st_mtime,
                reverse=True,
            )
            if not files:
                st.info("Chưa có file báo cáo để tải.")
                return
            for file_path in files:
                mime_type = mimetypes.guess_type(file_path.name)[0] or "application/octet-stream"
                with file_path.open("rb") as file_obj:
                    st.download_button(
                        label=f"⬇️ Tải {file_path.name}",
                        data=file_obj.read(),
                        file_name=file_path.name,
                        mime=mime_type,
                        key=f"cloud_download_{file_path.name}",
                    )
        except Exception as exc:
            import streamlit as st
            st.error(f"Không thể tạo khu vực tải file: {exc}")

    os.startfile = _startfile_linux_compat


# ---------------------------------------------------------------------------
# Bản tin mùa: hậu xử lý Bảng 1 để khớp mẫu người dùng cung cấp.
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Pt, Twips
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import bulletin.bulletin_generator as _bg

    _original_seasonal_bulletin = _bg.create_qtri_seasonal_bulletin

    def _format_seasonal_t1_cell(cell, text, bold=False, size=13):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.clear()
        r = p.add_run(str(text))
        r.bold = bold
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        _bg.set_cell_margins(cell, top=45, bottom=45, left=35, right=35)

    def _replace_seasonal_table1(output_path, data_dict):
        doc = Document(output_path)
        old_table = None
        for table in doc.tables:
            txt = " ".join(c.text for row in table.rows for c in row.cells)
            if "Yếu tố" in txt and ("Hmax" in txt or "Hmin" in txt):
                old_table = table
                break
        if old_table is None:
            return

        labels = list(data_dict.get("table1_labels", []))[:3]
        rows = list(data_dict.get("table1_rows", []))[:3]
        while len(labels) < 3:
            labels.append("-")
        while len(rows) < 3:
            rows.append({})
        station_name = data_dict.get("table1_station_name", "Cồn Cỏ")

        new_table = doc.add_table(rows=1, cols=6)
        new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        new_table.autofit = False
        _bg.set_table_borders(new_table)

        hdr = new_table.rows[0].cells
        merged_hdr = hdr[0].merge(hdr[2])
        _format_seasonal_t1_cell(merged_hdr, "Yếu tố", bold=True, size=13)
        for i, label in enumerate(labels):
            _format_seasonal_t1_cell(hdr[3 + i], label, bold=True, size=13)

        specs = [
            ("Thủy triều", "Nước lớn", "Hmax (cm)", "hmax"),
            ("Thủy triều", "Nước lớn", "Ngày xuất hiện", "hmax_day"),
            ("Thủy triều", "Nước ròng", "Hmin (cm)", "hmin"),
            ("Thủy triều", "Nước ròng", "Ngày xuất hiện", "hmin_day"),
            ("Sóng", "", "Độ cao sóng lớn nhất (m)", "wave_max"),
            ("Sóng", "", "Hướng sóng", "wave_dir"),
            ("Sóng", "", "Ngày xuất hiện", "wave_day"),
        ]

        group_starts = {}
        group_ends = {}
        for group_name, sub_name, metric, key in specs:
            cells = new_table.add_row().cells
            row_idx = len(new_table.rows) - 1
            group_starts.setdefault(group_name, row_idx)
            group_ends[group_name] = row_idx
            _format_seasonal_t1_cell(cells[1], sub_name, size=13)
            _format_seasonal_t1_cell(cells[2], metric, size=13)
            for period_idx, ext in enumerate(rows):
                value = ext.get(key) if ext else None
                if value is None or value == "":
                    value = "-"
                elif key in ("hmax", "hmin"):
                    value = str(int(round(float(value))))
                elif key in ("hmax_day", "hmin_day", "wave_day"):
                    value = str(int(value)) if str(value).isdigit() else str(value)
                _format_seasonal_t1_cell(cells[3 + period_idx], value, size=13)

        for group_name in ("Thủy triều", "Sóng"):
            _bg.merge_vertical(
                new_table, 0, group_starts[group_name], group_ends[group_name],
                group_name, bold=False, size=13,
            )

        widths = [1450, 1600, 1850, 1390, 1390, 1390]
        for i, width in enumerate(widths):
            new_table.columns[i].width = Twips(width)
            for row in new_table.rows:
                row.cells[i].width = Twips(width)

        old_element = old_table._element
        parent = old_element.getparent()
        parent.replace(old_element, new_table._element)

        for paragraph in doc.paragraphs:
            if paragraph.text.strip().startswith("Bảng 1:"):
                first_month = last_day = last_month = last_year = None
                m0 = re.search(r"(?:Tháng|tháng)\s+(\d+)/(\d+)", labels[0])
                m2 = re.search(r"01-(\d+)/(\d+)/(\d+)", labels[2])
                if m0:
                    first_month = int(m0.group(1))
                if m2:
                    last_day = int(m2.group(1))
                    last_month = int(m2.group(2))
                    last_year = int(m2.group(3))
                title = paragraph.text
                if first_month is not None and last_day is not None:
                    title = (
                        f"Bảng 1: Đặc trưng sóng, thủy triều tại trạm Hải văn {station_name} "
                        f"từ tháng {first_month} đến ngày {last_day} tháng {last_month}/{last_year}"
                    )
                paragraph.clear()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(2)
                r = paragraph.add_run(title)
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(13)
                break
        doc.save(output_path)

    def _fixed_seasonal_bulletin(data_dict, forecaster=None, issue_time=None,
                                 output_path="Ban_tin_Hai_van_Mua_Quang_Tri.docx"):
        result = _original_seasonal_bulletin(
            data_dict, forecaster=forecaster, issue_time=issue_time, output_path=output_path
        )
        try:
            _replace_seasonal_table1(result, data_dict)
        except Exception:
            pass
        return result

    _bg.create_qtri_seasonal_bulletin = _fixed_seasonal_bulletin

except Exception:
    pass


# ---------------------------------------------------------------------------
# HỒ SƠ MÙA: PHỤ LỤC 1 + PHỤ LỤC 2 PHẢI DÙNG MÉT (m).
#
# Bản generator gốc của dự án từng dựng hai phụ lục với use_cm=True. Vì
# app.py import create_forecast_dossier trực tiếp, ta hậu xử lý chính file
# DOCX sau khi sinh để bảo đảm cả tiêu đề đơn vị và trị số đều được đổi từ
# cm sang m, bất kể đường gọi nào tạo hồ sơ.
# ---------------------------------------------------------------------------
try:
    import bulletin.bulletin_generator as _bg2
    import bulletin.dossier_generator as _dg

    # Trước hết vẫn ép đường dựng phụ lục dùng mét ngay từ nguồn.
    def _fixed_seasonal_appendix_tables(doc, data_dict):
        section = doc.sections[-1]
        n_months = len(data_dict.get("forecast_months", [])) or 3
        col_widths = _bg2._compute_tide_table_col_widths(
            section.page_width, section.left_margin, section.right_margin, n_months
        )
        _bg2._render_tide_zone_table(
            doc, data_dict,
            "Phụ lục 1: Kết quả dự báo theo phương pháp phân tích hàm điều hòa",
            col_widths, use_cm=False,
        )
        _bg2._render_tide_zone_table(
            doc, data_dict,
            "Phụ lục 2: Chọn kết quả dự báo",
            col_widths, use_cm=False,
        )

    _bg2._add_seasonal_appendix_tables = _fixed_seasonal_appendix_tables

    _original_create_dossier = _dg.create_forecast_dossier

    def _replace_text_in_cell(cell, old, new):
        for paragraph in cell.paragraphs:
            if old in paragraph.text:
                for run in paragraph.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)

    def _convert_appendix_cm_to_m(output_path):
        """Bảo đảm Phụ lục 1/2 trong hồ sơ mùa hiển thị mét và quy đổi số."""
        doc = Document(output_path)
        changed = False

        for table in doc.tables:
            # Tìm đúng bảng có các cột Hx/Hm. Không đụng các bảng nghiệp vụ khác.
            texts = [
                (cell.text or "").strip()
                for row in table.rows
                for cell in row.cells
            ]
            joined = " | ".join(texts)
            if "Hx (cm)" not in joined and "Hm (cm)" not in joined:
                continue

            hx_cols = set()
            hm_cols = set()
            for row in table.rows[:4]:
                for idx, cell in enumerate(row.cells):
                    txt = (cell.text or "").strip()
                    if "Hx (cm)" in txt:
                        hx_cols.add(idx)
                    if "Hm (cm)" in txt:
                        hm_cols.add(idx)

            # Đổi tiêu đề đơn vị.
            for row in table.rows:
                for cell in row.cells:
                    _replace_text_in_cell(cell, "Hx (cm)", "Hx (m)")
                    _replace_text_in_cell(cell, "Hm (cm)", "Hm (m)")
                    _replace_text_in_cell(cell, "(cm)", "(m)")

            # Quy đổi chỉ các cột Hx/Hm; không đụng cột Ngày/Thời gian.
            for row_idx, row in enumerate(table.rows):
                if row_idx < 3:
                    continue
                for col_idx in hx_cols | hm_cols:
                    if col_idx >= len(row.cells):
                        continue
                    cell = row.cells[col_idx]
                    raw = (cell.text or "").strip().replace(",", ".")
                    if not re.fullmatch(r"-?\d+(?:\.\d+)?", raw):
                        continue
                    try:
                        value_m = float(raw) / 100.0
                    except ValueError:
                        continue
                    # Viết theo đúng đơn vị mét, 2 chữ số thập phân.
                    p = cell.paragraphs[0]
                    p.clear()
                    r = p.add_run(f"{value_m:.2f}")
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    changed = True

        if changed:
            doc.save(output_path)

    def _fixed_create_dossier(data_dict, output_path="Ho_so_du_bao_Quang_Tri.docx"):
        result = _original_create_dossier(data_dict, output_path=output_path)
        if isinstance(data_dict, dict) and data_dict.get("zone_table_kind") == "seasonal_appendix":
            try:
                _convert_appendix_cm_to_m(result)
            except Exception:
                pass
        return result

    _dg.create_forecast_dossier = _fixed_create_dossier

except Exception:
    pass
